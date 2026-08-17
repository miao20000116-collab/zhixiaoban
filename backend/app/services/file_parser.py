"""JD / document text extraction utilities."""

from pathlib import Path


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".doc":
        raise ValueError(
            "暂不支持旧版 Word（.doc）。请另存为 .docx 或导出 PDF 后再上传。"
        )

    if suffix in {".txt", ".md", ".markdown", ".csv"}:
        return content.decode("utf-8", errors="ignore").strip()

    if suffix == ".pdf":
        return _extract_pdf(content)

    if suffix in {".docx"}:
        return _extract_docx(content)

    # Fallback: try utf-8 decode
    try:
        text = content.decode("utf-8")
        if text.strip():
            return text.strip()
    except UnicodeDecodeError:
        pass

    raise ValueError(
        f"不支持的文件类型: {suffix or 'unknown'}。支持：.txt .md .pdf .docx"
    )


def _extract_pdf(content: bytes) -> str:
    try:
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("PDF 中没有可提取的文字（可能是扫描件图片）")
        return text
    except ImportError as exc:
        raise ValueError("PDF 支持需要安装 pypdf：pip install pypdf") from exc


def _extract_docx(content: bytes) -> str:
    try:
        from io import BytesIO

        from docx import Document

        doc = Document(BytesIO(content))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()
            if line:
                parts.append(line)

        # Many Chinese resumes put key content in tables — include them.
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    # Deduplicate merged-cell repeats while keeping order
                    deduped: list[str] = []
                    for cell in cells:
                        if not deduped or deduped[-1] != cell:
                            deduped.append(cell)
                    parts.append(" | ".join(deduped))

        text = "\n".join(parts).strip()
        if not text:
            raise ValueError("Word 文档中没有可提取的文字")
        return text
    except ImportError as exc:
        raise ValueError("DOCX 支持需要安装 python-docx：pip install python-docx") from exc
